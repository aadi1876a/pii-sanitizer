import os
import pandas as pd
import json
from docx import Document
from PIL import Image, ImageDraw


class FileSanitizer:

    def mask_email(self, value):
        try:
            name, domain = value.split("@")
            return name[0] + "***@" + domain
        except:
            return "[MASKED]"

    def mask_phone(self, value):
        if len(value) >= 4:
            return value[:2] + "******" + value[-2:]
        return "[MASKED]"

    def mask_dob(self, value):
        return "**/**/****"

    def redact(self):
        return "[REDACTED]"

    def visual_mask(self, value):
        return "█" * len(value)

    # -------------------------------------------------
    # RULE MAPPING
    # -------------------------------------------------

    def apply_rule(self, pii_type, value):

        # MASK
        if pii_type == "email":
            return self.mask_email(value)

        if pii_type == "phone":
            return self.mask_phone(value)

        if pii_type == "dob":
            return self.mask_dob(value)

        # REDACT
        if pii_type in [
            "pan",
            "aadhaar",
            "ip_address",       # redacted as requested
            "card_number",
            "upi_id",
            "ifsc",
            "cvv",
            "bank_account",     # redacted as requested
        ]:
            return self.redact()

        # VISUAL MASK
        if pii_type in [
            "passport",
            "vehicle_reg",
            "voter_id",
            "location",
            "address",          # visual mask (house number already kept by detector)
            "device_id",        # visual mask as requested
        ]:
            return self.visual_mask(value)

        # DO NOT CHANGE
        if pii_type in ["name", "organisation", "url"]:
            return value

        return value

    # -------------------------------------------------
    # REMOVE OVERLAPPING DETECTIONS
    # -------------------------------------------------

    def remove_overlaps(self, findings):

        priority = {
            "pan": 1,
            "aadhaar": 1,
            "ip_address": 1,
            "card_number": 1,
            "upi_id": 1,
            "ifsc": 1,
            "email": 1,
            "phone": 1,
            "dob": 1,
            "bank_account": 1,
            "address": 1,
            "device_id": 1,
        }

        findings = sorted(
            findings,
            key=lambda x: (x["start"], priority.get(x["type"], 2))
        )

        filtered = []
        last_end = -1

        for f in findings:
            if f["start"] >= last_end:
                filtered.append(f)
                last_end = f["end"]

        return filtered

    # -------------------------------------------------
    # FILTER OUT FIELD LABELS
    # -------------------------------------------------

    def filter_labels(self, text, findings):
        result = []
        for f in findings:
            end = f["end"]
            lookahead = text[end: end + 2].strip()
            if not lookahead.startswith(":"):
                result.append(f)
        return result

    # -------------------------------------------------
    # POSITION BASED SANITIZATION
    # -------------------------------------------------

    def sanitize_with_positions(self, text, findings):

        if not findings:
            return text

        findings = self.filter_labels(text, findings)
        findings = self.remove_overlaps(findings)

        sanitized_parts = []
        last_index = 0

        for item in findings:
            start = item["start"]
            end = item["end"]
            value = item["value"]
            pii_type = item["type"]

            sanitized_parts.append(text[last_index:start])
            replacement = self.apply_rule(pii_type, value)
            sanitized_parts.append(replacement)
            last_index = end

        sanitized_parts.append(text[last_index:])
        return "".join(sanitized_parts)

    def sanitize_text(self, text, findings):
        if not findings:
            return text
        return self.sanitize_with_positions(text, findings)

    def sanitize_cell(self, cell_text, detector):
        if not cell_text or not cell_text.strip():
            return cell_text
        cell_findings = detector.detect_with_positions(cell_text)
        return self.sanitize_text(cell_text, cell_findings)

    # -------------------------------------------------
    # FILE TYPE HANDLERS
    # -------------------------------------------------

    def sanitize_txt(self, input_path, output_path, findings):
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()
        clean = self.sanitize_text(text, findings)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(clean)

    def sanitize_csv(self, input_path, output_path, findings, detector=None):
        df = pd.read_csv(input_path)
        if detector is not None:
            for col in df.columns:
                df[col] = df[col].astype(str).apply(
                    lambda x: self.sanitize_cell(x, detector)
                )
        else:
            for col in df.columns:
                df[col] = df[col].astype(str).apply(
                    lambda x: self.sanitize_text(x, findings)
                )
        df.to_csv(output_path, index=False)

    def sanitize_json(self, input_path, output_path, findings):
        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        def sanitize_value(v):
            if isinstance(v, str):
                return self.sanitize_text(v, findings)
            if isinstance(v, list):
                return [sanitize_value(i) for i in v]
            if isinstance(v, dict):
                return {k: sanitize_value(val) for k, val in v.items()}
            return v

        clean = sanitize_value(data)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(clean, f, indent=4)

    def sanitize_docx(self, input_path, output_path, findings, detector=None):
        doc = Document(input_path)

        def sanitize_paragraph(para):
            # Build the full paragraph text and a character-to-run map
            full_text = ""
            char_map = []  # each index → (run_index, char_index_in_run)
            for r_idx, run in enumerate(para.runs):
                for c_idx, ch in enumerate(run.text):
                    char_map.append((r_idx, c_idx))
                full_text += run.text

            if not full_text.strip():
                return

            # Detect PII on the full paragraph text
            if detector is not None:
                para_findings = detector.detect_with_positions(full_text)
                para_findings = self.filter_labels(full_text, para_findings)
                para_findings = self.remove_overlaps(para_findings)
            else:
                para_findings = self.filter_labels(full_text, findings)
                para_findings = self.remove_overlaps(findings)

            if not para_findings:
                return

            # Build sanitized full text
            sanitized = self.sanitize_text(full_text, para_findings)

            # If nothing changed, skip
            if sanitized == full_text:
                return

            # Build list of (start, end, replacement) from findings
            replacements = []
            for item in para_findings:
                replacement = self.apply_rule(item["type"], item["value"])
                replacements.append((item["start"], item["end"], replacement))

            # Walk through original text and assign new chars to runs
            # preserving run boundaries (and thus formatting) wherever possible
            run_texts = [""] * len(para.runs)
            pos = 0
            rep_idx = 0

            while pos < len(full_text):
                # Check if we're at a replacement start
                if rep_idx < len(replacements) and pos == replacements[rep_idx][0]:
                    r_start, r_end, repl = replacements[rep_idx]
                    # Assign the replacement text to the run that owns r_start
                    run_owner = char_map[r_start][0] if r_start < len(char_map) else len(para.runs) - 1
                    run_texts[run_owner] += repl
                    pos = r_end
                    rep_idx += 1
                else:
                    # Normal character — assign to its original run
                    if pos < len(char_map):
                        run_owner = char_map[pos][0]
                        run_texts[run_owner] += full_text[pos]
                    pos += 1

            # Write the new texts back into runs (bold/italic/font preserved)
            for r_idx, run in enumerate(para.runs):
                run.text = run_texts[r_idx]

        # Process all paragraphs
        for para in doc.paragraphs:
            sanitize_paragraph(para)

        # Also process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        sanitize_paragraph(para)

        doc.save(output_path)

    def sanitize_sql(self, input_path, output_path, findings):
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()
        clean = self.sanitize_text(text, findings)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(clean)

    def sanitize_pdf(self, input_path, output_path, findings, detector=None):
        try:
            import fitz  # PyMuPDF — install with: pip install pymupdf

            doc = fitz.open(input_path)

            for page in doc:
                blocks = page.get_text("dict")["blocks"]

                for block in blocks:
                    if block.get("type") != 0:  # 0 = text block
                        continue

                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            original_text = span.get("text", "")
                            if not original_text.strip():
                                continue

                            # Sanitize this span text
                            if detector is not None:
                                span_findings = detector.detect_with_positions(original_text)

                                # Skip field labels (token immediately followed by ":")
                                valid = []
                                for f in span_findings:
                                    lookahead = original_text[f["end"]: f["end"] + 2].strip()
                                    if lookahead.startswith(":"):
                                        continue
                                    valid.append(f)

                                # Remove overlaps — specific PII beats bank_account
                                priority = {
                                    "phone": 1, "email": 1, "pan": 1, "aadhaar": 1,
                                    "card_number": 1, "upi_id": 1, "ifsc": 1, "dob": 1,
                                    "ip_address": 1, "cvv": 1, "passport": 1,
                                    "voter_id": 1, "vehicle_reg": 1, "device_id": 1,
                                    "bank_account": 2,
                                }
                                valid.sort(key=lambda x: (x["start"], priority.get(x["type"], 2)))
                                deduped, last_end = [], -1
                                for f in valid:
                                    if f["start"] >= last_end:
                                        deduped.append(f)
                                        last_end = f["end"]

                                sanitized_text = self.sanitize_text(original_text, deduped)
                            else:
                                sanitized_text = self.sanitize_text(original_text, findings)

                            if sanitized_text == original_text:
                                continue  # nothing changed, skip

                            bbox = fitz.Rect(span["bbox"])
                            font_size = span.get("size", 10)
                            font_color = span.get("color", 0)

                            # Convert int color → (r, g, b) floats
                            r = ((font_color >> 16) & 0xFF) / 255.0
                            g = ((font_color >> 8) & 0xFF) / 255.0
                            b = (font_color & 0xFF) / 255.0

                            # Pick font based on bold/italic flags
                            flags = span.get("flags", 0)
                            is_bold   = bool(flags & (1 << 4))
                            is_italic = bool(flags & (1 << 1))
                            if is_bold and is_italic:
                                fontname = "helv"
                            elif is_bold:
                                fontname = "hebo"
                            elif is_italic:
                                fontname = "heit"
                            else:
                                fontname = "helv"

                            # Erase original text with white rectangle
                            page.draw_rect(bbox, color=None, fill=(1, 1, 1))

                            # insert_text origin is the BASELINE (bottom-left of text),
                            # not the top-left. Use bbox.y1 minus a small descender gap
                            # so the text sits vertically centred inside the original bbox.
                            baseline = fitz.Point(bbox.x0, bbox.y1 - (bbox.height - font_size) / 2 - 1)

                            page.insert_text(
                                baseline,
                                sanitized_text,
                                fontsize=font_size,
                                fontname=fontname,
                                color=(r, g, b),
                            )

            doc.save(output_path, garbage=4, deflate=True)
            doc.close()

        except ImportError:
            raise Exception(
                "PyMuPDF is required for format-preserving PDF sanitization. "
                "Install it with: pip install pymupdf"
            )
        except Exception as e:
            raise Exception(f"PDF sanitization failed: {str(e)}")

    def sanitize_image(self, input_path, output_path, detector=None):
        """
        OCR-based image sanitization.
        - Uses pytesseract to extract word-level bounding boxes.
        - Reconstructs line-level text, runs PII detector on each line.
        - Draws a filled black rectangle precisely over each PII word/span.
        Falls back to a centre-strip redaction if pytesseract is not installed.
        """
        img = Image.open(input_path).convert("RGB")

        try:
            import pytesseract

            # ------------------------------------------------------------------
            # 1. Get word-level OCR data (text + bbox per word)
            # ------------------------------------------------------------------
            data = pytesseract.image_to_data(
                img, output_type=pytesseract.Output.DICT
            )

            n = len(data["text"])

            # Group words into lines using (block_num, par_num, line_num)
            lines = {}
            for i in range(n):
                word = data["text"][i].strip()
                if not word:
                    continue
                key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
                lines.setdefault(key, []).append(i)

            draw = ImageDraw.Draw(img)

            # ------------------------------------------------------------------
            # 2. For each line, build full text and detect PII positions
            # ------------------------------------------------------------------
            for key, indices in lines.items():
                # Reconstruct line text with spaces between words
                words = [data["text"][i] for i in indices]
                line_text = " ".join(words)

                if not line_text.strip():
                    continue

                # Use provided detector or fall back to simple pattern scan
                if detector is not None:
                    line_findings = detector.detect_with_positions(line_text)
                    line_findings = self.filter_labels(line_text, line_findings)
                    line_findings = self.remove_overlaps(line_findings)
                else:
                    continue  # no detector → nothing to redact precisely

                if not line_findings:
                    continue

                # -------------------------------------------------------
                # 3. Map character positions back to word bounding boxes
                # -------------------------------------------------------
                # Build char→word_index map for this line
                char_to_word = []
                for w_idx, word in enumerate(words):
                    for _ in word:
                        char_to_word.append(w_idx)
                    if w_idx < len(words) - 1:
                        char_to_word.append(w_idx)  # the space between words

                def word_bbox(w_idx):
                    i = indices[w_idx]
                    x = data["left"][i]
                    y = data["top"][i]
                    return (x, y, x + data["width"][i], y + data["height"][i])

                # Calibrate font size by matching the rendered WIDTH of a
                # non-PII label word (e.g. "Phone:") against its OCR pixel width.
                # Width is far more reliable than height for font size matching.
                from PIL import ImageFont as _IFont
                FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

                pii_values = {f["value"] for f in line_findings}
                ref_idx = next(
                    (i for i in indices
                     if data["text"][i].strip()
                     and data["text"][i].strip() not in pii_values
                     and len(data["text"][i].strip()) >= 3),
                    None
                )

                if ref_idx is not None:
                    ref_word   = data["text"][ref_idx].strip()
                    ref_width  = data["width"][ref_idx]
                    ref_height = data["height"][ref_idx]

                    def _word_w(pt):
                        try:
                            f = _IFont.truetype(FONT_PATH, pt)
                            bb = f.getbbox(ref_word)
                            return bb[2] - bb[0]
                        except Exception:
                            return pt

                    # Width-based calibration via binary search
                    lo, hi = 6, 300
                    while lo < hi:
                        mid = (lo + hi) // 2
                        if _word_w(mid) < ref_width:
                            lo = mid + 1
                        else:
                            hi = mid
                    width_font_size = lo

                    # Height-based calibration: find pt where rendered height
                    # matches the OCR bounding-box height of the reference word.
                    def _word_h(pt):
                        try:
                            f = _IFont.truetype(FONT_PATH, pt)
                            bb = f.getbbox(ref_word)
                            return bb[3] - bb[1]
                        except Exception:
                            return pt

                    lo2, hi2 = 6, 300
                    while lo2 < hi2:
                        mid2 = (lo2 + hi2) // 2
                        if _word_h(mid2) < ref_height:
                            lo2 = mid2 + 1
                        else:
                            hi2 = mid2
                    height_font_size = lo2

                    # Use the larger estimate so the replacement is never
                    # visually smaller than the original text.
                    line_font_size = max(width_font_size, height_font_size)
                else:
                    # Fallback: use OCR height of tallest word
                    ref_heights = [data["height"][i] for i in indices
                                   if data["text"][i].strip() and data["height"][i] >= 20]
                    line_font_size = max(ref_heights) if ref_heights else 56

                for finding in line_findings:
                    start_char = finding["start"]
                    end_char   = finding["end"] - 1

                    # Clamp to valid range
                    start_char = min(start_char, len(char_to_word) - 1)
                    end_char   = min(end_char,   len(char_to_word) - 1)

                    if start_char < 0 or end_char < 0:
                        continue

                    first_word = char_to_word[start_char]
                    last_word  = char_to_word[end_char]

                    # Union bbox of all words in the PII span
                    x0 = min(word_bbox(w)[0] for w in range(first_word, last_word + 1))
                    y0 = min(word_bbox(w)[1] for w in range(first_word, last_word + 1))
                    x1 = max(word_bbox(w)[2] for w in range(first_word, last_word + 1))
                    y1 = max(word_bbox(w)[3] for w in range(first_word, last_word + 1))

                    padding = max(4, int(line_font_size * 0.1))

                    # Apply the same rule as all other file types
                    replacement = self.apply_rule(finding["type"], finding["value"])

                    # DO NOT CHANGE → leave the image pixels untouched
                    if replacement == finding["value"]:
                        continue

                    # Erase original text with a white rectangle
                    draw.rectangle([(x0, y0 - padding), (x1, y1 + padding)], fill="white")

                    # line_font_size is already the calibrated point size
                    # (matched via label word width), so use it directly.
                    from PIL import ImageFont
                    FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
                    try:
                        font = ImageFont.truetype(FONT_PATH, line_font_size)
                    except Exception:
                        font = ImageFont.load_default()

                    # Align ink top to y0
                    try:
                        ink_top = font.getbbox("Eg")[1]
                    except Exception:
                        ink_top = 0

                    draw.text(
                        (x0, y0 - ink_top),
                        replacement,
                        fill=(0, 0, 0),
                        font=font,
                    )

        except ImportError:
            # Graceful fallback — centre-strip redaction
            draw = ImageDraw.Draw(img)
            w, h = img.size
            draw.rectangle([(w * 0.3, h * 0.3), (w * 0.7, h * 0.5)], fill="black")

        img.save(output_path)

    # -------------------------------------------------
    # MAIN ROUTER
    # -------------------------------------------------

    def sanitize_file(self, input_path, output_folder, findings, detector=None):

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        output_path = os.path.join(output_folder, name + "_sanitized" + ext)
        ext = ext.lower()

        if ext == ".txt":
            self.sanitize_txt(input_path, output_path, findings)
        elif ext == ".csv":
            self.sanitize_csv(input_path, output_path, findings, detector)
        elif ext == ".json":
            self.sanitize_json(input_path, output_path, findings)
        elif ext == ".docx":
            self.sanitize_docx(input_path, output_path, findings, detector)
        elif ext == ".sql":
            self.sanitize_sql(input_path, output_path, findings)
        elif ext == ".pdf":
            self.sanitize_pdf(input_path, output_path, findings, detector)
        elif ext in [".png", ".jpg", ".jpeg"]:
            self.sanitize_image(input_path, output_path, detector=detector)
        else:
            raise Exception("Unsupported file type")

        return output_path


if __name__ == "__main__":

    from backend.detector.detector import PIIDetector

    input_file = os.path.join("inputs", "sample.txt")
    output_folder = "outputs"

    sanitizer = FileSanitizer()
    detector = PIIDetector()

    with open(input_file, "r", encoding="utf-8") as f:
        text = f.read()

    findings = detector.detect_with_positions(text)
    output_path = sanitizer.sanitize_file(input_file, output_folder, findings, detector=detector)
    print("Sanitized file created at:", output_path)